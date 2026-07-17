import streamlit as st
import io
import csv
import requests
from docx import Document
import pdfplumber

# Импорты для генерации PDF через ReportLab
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Инициализация конфигурации страницы Streamlit (ДОЛЖНА БЫТЬ В САМОМ НАЧАЛЕ)
st.set_page_config(
    page_title="Переводчик Сертификатов Анализа (CoA)",
    page_icon="🧪",
    layout="wide"
)

# ==========================================
# 0. РЕГИСТРАЦИЯ ШРИФТОВ CALIBRI
# ==========================================
try:
    pdfmetrics.registerFont(TTFont('Calibri', 'CALIBRI.TTF'))
    pdfmetrics.registerFont(TTFont('Calibri-Bold', 'CALIBRIB.TTF'))
    pdfmetrics.registerFont(TTFont('Calibri-Italic', 'CALIBRII.TTF'))
except Exception as e:
    st.error(f"Ошибка регистрации шрифтов Calibri: {e}. Убедитесь, что файлы CALIBRI.TTF, CALIBRIB.TTF, CALIBRII.TTF загружены в корень проекта.")

# ==========================================
# 1. ЗАГРУЗКА И ИНИЦИАЛИЗАЦИЯ СЛОВАРЯ (SESSION STATE)
# ==========================================

def load_dictionary_from_csv(file_path="dictionary.csv"):
    """Читает CSV словарь из репозитория и переводит в чистый dict"""
    coa_dict = {}
    try:
        with open(file_path, mode='r', encoding='utf-8') as f:
            reader = csv.reader(f, delimiter=';')
            next(reader)  # Пропускаем заголовок english;russian
            for row in reader:
                if len(row) >= 2:
                    eng_phrase = row[0].strip().lower()
                    rus_phrase = row[1].strip()
                    coa_dict[eng_phrase] = rus_phrase
    except Exception as e:
        st.error(f"Не удалось загрузить dictionary.csv: {e}")
    return coa_dict

# Храним словарь в session_state, чтобы изменения не стирались при перезапуске страницы
if "coa_dictionary" not in st.session_state:
    st.session_state["coa_dictionary"] = load_dictionary_from_csv()

# ==========================================
# 2. МОДУЛЬ ПАРСИНГА (DOCX & PDF)
# ==========================================
def parse_docx(file_bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    full_text = []
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for paragraph in doc.paragraphs:
                if paragraph._element == element:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                    break
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    for row in table.rows:
                        row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        full_text.append(" | ".join(row_text))
                    break
    return "\n".join(full_text)

def parse_pdf(file_bytes) -> str:
    full_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            page_text = page.extract_text()
            if page_text:
                lines = page_text.split("\n")
                for line in lines:
                    full_text.append(line)
            if tables:
                full_text.append("\n--- Обнаружена таблица показателей ---")
                for table in tables:
                    for row in table:
                        row_text = [str(cell).strip().replace("\n", " ") if cell else "" for cell in row]
                        full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def extract_text_from_file(uploaded_file) -> str:
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".docx"):
        return parse_docx(file_bytes)
    elif file_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif file_name.endswith(".txt"):
        return file_bytes.decode("utf-8")
    else:
        raise ValueError("Неподдерживаемый формат файла")

# ==========================================
# 3. МОДУЛЬ ГИБРИДНОГО ПЕРЕВОДА И АВТОПОПОЛНЕНИЯ
# ==========================================
import base64

def add_to_github_dictionary(eng_phrase: str, rus_phrase: str):
    """Автоматически добавляет новую пару слов в dictionary.csv на GitHub"""
    if "github" not in st.secrets:
        return
        
    g_sec = st.secrets["github"]
    token = g_sec.get("token")
    repo = g_sec.get("repo")
    branch = g_sec.get("branch", "main")
    file_path = "dictionary.csv"
    
    if not token or not repo:
        return

    url = f"https://api.github.com/repos/{repo}/contents/{file_path}"
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json"
    }
    
    try:
        response = requests.get(url, headers=headers, params={"ref": branch}, timeout=10)
        
        if response.status_code == 200:
            file_data = response.json()
            sha = file_data["sha"]
            current_content = base64.b64decode(file_data["content"]).decode("utf-8")
        elif response.status_code == 404:
            sha = None
            current_content = "english;russian"
        else:
            return # Ошибка доступа к GitHub API
        new_line = f"{eng_phrase.strip().lower()};{rus_phrase.strip()}"
        if new_line in current_content:
            return

        if current_content.endswith("\n") or current_content.endswith("\r"):
            updated_content = current_content + new_line + "\n"
        else:
            updated_content = current_content + "\n" + new_line + "\n"

        encoded_content = base64.b64encode(updated_content.encode("utf-8")).decode("utf-8")
        
        # Шаг 3: Отправляем коммит на GitHub
        commit_data = {
            "message": f"🤖 Авто-добавление термина: {eng_phrase}",
            "content": encoded_content,
            "branch": branch
        }
        if sha:
            commit_data["sha"] = sha
            
        put_response = requests.put(url, headers=headers, json=commit_data, timeout=10)
        
        if put_response.status_code in [200, 201]:
            # Важнейший шаг: пишем в session_state, чтобы изменение применилось мгновенно!
            st.session_state["coa_dictionary"][eng_phrase.strip().lower()] = rus_phrase.strip()
    except Exception:
        pass

def save_full_dictionary_to_github(dataframe):
    """Полностью перезаписывает dictionary.csv на GitHub на основе измененного DataFrame"""
    if "github" not in st.secrets:
        st.error("Ключи GitHub не найдены в secrets!")
        return False
        
    g_sec = st.secrets["github"]
    token = g_sec.get("token")
    repo = g_sec.get("repo")
    branch = g_sec.get("branch", "main")
    file_path = "dictionary.csv"
    
    url = f"https://api.github.com/repos/{repo}/contents/{file_path}"
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json"
    }
    
    try:
        # Шаг 1: Конвертируем DataFrame в строку CSV с разделителем ';'
        # Убираем индекс pandas, чтобы он не попадал в файл
        csv_buffer = io.StringIO()
        dataframe.to_csv(csv_buffer, sep=';', index=False, encoding='utf-8')
        new_csv_content = csv_buffer.getvalue()
        
        # Шаг 2: Получаем текущий SHA файла, чтобы разрешить перезапись
        response = requests.get(url, headers=headers, params={"ref": branch}, timeout=10)
        sha = response.json()["sha"] if response.status_code == 200 else None
        
        # Шаг 3: Кодируем новый контент в Base64
        encoded_content = base64.b64encode(new_csv_content.encode("utf-8")).decode("utf-8")
        
        # Шаг 4: Отправляем коммит
        commit_data = {
            "message": "📝 Ручное обновление словаря через интерфейс Streamlit",
            "content": encoded_content,
            "branch": branch
        }
        if sha:
            commit_data["sha"] = sha
            
        put_response = requests.put(url, headers=headers, json=commit_data, timeout=10)
        
        if put_response.status_code in [200, 201]:
            # Перечитываем обновленный словарь напрямую в сессию
            st.session_state["coa_dictionary"] = load_dictionary_from_csv()
            return True
        else:
            st.error(f"GitHub API вернул ошибку: {put_response.status_code}")
            return False
    except Exception as e:
        st.error(f"Не удалось сохранить словарь на GitHub: {e}")
        return False

def translate_via_yandex_gpt(text: str) -> str:
    if not text.strip():
        return text
    if "yandex" not in st.secrets:
        return f"[Ключи Yandex не найдены] {text}"
        
    folder_id = st.secrets["yandex"]["folder_id"]
    api_key = st.secrets["yandex"]["api_key"]
    
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {"Authorization": f"Api-Key {api_key}", "Content-Type": "application/json"}
    
    data = {
        "modelUri": f"gpt://{folder_id}/yandexgpt-lite",
        "completionOptions": {"stream": False, "temperature": 0.1, "maxTokens": "2000"},
        "messages": [
            {
                "role": "system",
                "text": "Ты профессиональный переводчик химической документации. Переведи строго текст на русский язык. Сохраняй химические формулы, спецсимволы и цифры неизменными. Не добавляй отсебятины."
            },
            {"role": "user", "text": text}
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=15)
        if response.status_code == 200:
            return response.json()["result"]["alternatives"][0]["message"]["text"].strip()
        return f"[Ошибка API {response.status_code}] {text}"
    except Exception as e:
        return f"[Ошибка: {e}] {text}"


def translate_cell_or_line(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return ""
        
    # 1. Проверка на цифры/формулы
    if any(char.isdigit() for char in cleaned) and not any(char.isalpha() for char in cleaned):
        return cleaned
        
 # 2. Поиск в существующем CSV-словаре, хранящемся в session_state
    cleaned_lower = cleaned.lower()
    local_dict = st.session_state["coa_dictionary"]
    if cleaned_lower in local_dict:
        return local_dict[cleaned_lower]
        
    # 3. Если совпадений нет — переводим через YandexGPT
    translated = translate_via_yandex_gpt(cleaned)
    
    # 4. Если перевод успешный (не вернул ошибку API), отправляем новинку на GitHub
    if translated and not translated.startswith("[Ошибка"):
        # Запускаем фоновую отправку на GitHub
        add_to_github_dictionary(cleaned, translated)
        
    return translated


def process_coa_translation(raw_text: str, custom_name_ru: str = "") -> str:
    """Разбивает текст на строки, бережно переводит ячейки таблиц через '|' и собирает обратно"""
    translated_lines = []
    
    for line in raw_text.split("\n"):
        if "|" in line:
            cells = line.split("|")
            
            # 1. Проверяем маркеры строк
            is_product_name_row = any(marker in cells[0].lower() for marker in ["product name", "наименование продукта"])
            is_formula_row = any(marker in cells[0].lower() for marker in ["molecular formula", "молекулярная формула"])
            
            # 2. Логика для названия продукта (ручной ввод)
            if is_product_name_row and custom_name_ru.strip():
                translated_cells = [
                    translate_cell_or_line(cells[0]), 
                    custom_name_ru.strip()
                ]
                if len(cells) > 2:
                    translated_cells.extend([translate_cell_or_line(c) for c in cells[2:]])
                    
            # 3. Логика для молекулярной формулы (полная защита от цензуры нейросети)
            elif is_formula_row:
                translated_cells = [translate_cell_or_line(cells[0])] # Переводим только левую ячейку "MOLECULAR FORMULA"
                if len(cells) > 1:
                    translated_cells.append(cells[1].strip()) # Саму формулу (ячейку 1) оставляем как в оригинале
                if len(cells) > 2:
                    translated_cells.extend([translate_cell_or_line(c) for c in cells[2:]])
                    
            # 4. Обычные строки
            else:
                translated_cells = [translate_cell_or_line(cell) for cell in cells]
                
            translated_lines.append(" | ".join(translated_cells))
        else:
            translated_lines.append(translate_cell_or_line(line))
            
    return "\n".join(translated_lines)

# ==========================================
# 4. МОДУЛЬ ГЕНЕРАЦИИ PDF (REPORTLAB)
# ==========================================
def build_pdf_page_elements(story, text_content, is_russian=False):
    """Генерирует элементы (шапка, таблицы, дисклеймер, подпись) для одной языковой страницы"""
    
    # Стили текста
    style_site = ParagraphStyle('SiteStyle', fontName='Calibri-Bold', fontSize=14, alignment=TA_CENTER, textColor=colors.black)
    style_title = ParagraphStyle('TitleStyle', fontName='Calibri-Italic', fontSize=22, alignment=TA_CENTER, textColor=colors.HexColor('#0096D6'))
    style_cell = ParagraphStyle('CellStyle', fontName='Calibri', fontSize=10, textColor=colors.black)
    style_cell_header = ParagraphStyle('CellHeaderStyle', fontName='Calibri-Bold', fontSize=10, textColor=colors.white)
    style_disclaimer = ParagraphStyle('DisclaimerStyle', fontName='Calibri-Bold', fontSize=10, textColor=colors.black, leading=14)

    # 1. Логотип по центру
    try:
        story.append(Image("logo.jpg", width=160, height=50))
        story.append(Spacer(1, 10))
    except:
        story.append(Paragraph("[Логотип: logo.jpg не найден]", style_site))
        story.append(Spacer(1, 10))

    # 2. Веб-сайт
    story.append(Paragraph("www.spanlab.in", style_site))
    story.append(Spacer(1, 10))

    # 3. Заголовок листа
    title_text = "Сертификат анализа" if is_russian else "Certificate of Analysis"
    story.append(Paragraph(title_text, style_title))
    story.append(Spacer(1, 20))

    # Разбираем строки для Таблицы 1 и Таблицы 2
    lines = text_content.split("\n")
    table1_data = []
    table2_data = []
    disclaimer_text = []
    
    is_table2 = False
    is_disclaimer = False

    for line in lines:
        if not line.strip() or "---" in line or "spanlab.in" in line.lower() or "certificate of analysis" in line.lower() or "сертификат анализа" in line.lower():
            continue
            
        # Определяем начало Дисклеймера
        if "disclaimer" in line.lower() or "отказ от ответственности" in line.lower():
            is_disclaimer = True
            disclaimer_text.append(line)
            continue
        
        if is_disclaimer:
            disclaimer_text.append(line)
            continue

        # Определяем переключение на Таблицу 2 (Показатели)
        if "description" in line.lower() or "показатель" in line.lower():
            is_table2 = True

        if "|" in line:
            cells = [cell.strip() for cell in line.split("|")]
            if is_table2:
                table2_data.append(cells)
            else:
                table1_data.append(cells)

    # 4. Отрезаем и собираем Таблицу 1 (Метаданные продукта)
    if table1_data:
        # Оформляем текст в ячейках как Paragraph, чтобы он переносился
        formatted_table1 = []
        for r_idx, row in enumerate(table1_data):
            formatted_row = []
            for c_idx, cell in enumerate(row):
                # Переводим первую строку в UPPERCASE
                cell_txt = cell.upper() if r_idx == 0 else cell
                st_cell = style_cell_header if r_idx == 0 else style_cell
                formatted_row.append(Paragraph(cell_txt, st_cell))
            formatted_table1.append(formatted_row)

        t1 = Table(formatted_table1, colWidths=[200, 300])
        
        # Стиль для Таблицы 1 (Шапка голубая, строки чередуются)
        t1_styles = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0096D6')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ]
        # Добавляем чередование строк (зебру) для Таблицы 1
        for i in range(1, len(formatted_table1)):
            if i % 2 != 0:
                t1_styles.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F0F8FF'))) # Светло-голубой
        
        t1.setStyle(TableStyle(t1_styles))
        story.append(t1)

    # Разрыв строк (два переноса)
    story.append(Spacer(1, 20))

    # 5. Собираем Таблицу 2 (Показатели анализа)
    if table2_data:
        formatted_table2 = []
        for r_idx, row in enumerate(table2_data):
            formatted_row = []
            for c_idx, cell in enumerate(row):
                cell_txt = cell.upper() if r_idx == 0 else cell
                st_cell = style_cell_header if r_idx == 0 else style_cell
                formatted_row.append(Paragraph(cell_txt, st_cell))
            formatted_table2.append(formatted_row)

        # Вычисляем ширину колонок в зависимости от количества элементов (обычно 4 колонки)
        num_cols = len(table2_data[0]) if table2_data else 4
        col_widths = [150, 110, 110, 130] if num_cols == 4 else [500 / num_cols] * num_cols

        t2 = Table(formatted_table2, colWidths=col_widths)
        
        t2_styles = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0096D6')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ]
        # Зебра для Таблицы 2
        for i in range(1, len(formatted_table2)):
            if i % 2 != 0:
                t2_styles.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F0F8FF')))
                
        t2.setStyle(TableStyle(t2_styles))
        story.append(t2)

    story.append(Spacer(1, 20))

    # 6. Блок дисклеймера
    if disclaimer_text:
        full_disc_txt = " ".join(disclaimer_text)
        story.append(Paragraph(full_disc_txt, style_disclaimer))
        story.append(Spacer(1, 15))

    # 7. Вставка картинки подписи (sign.jpg) в самый конец
    try:
        story.append(Image("sign.jpg", width=120, height=60, hAlign='LEFT'))
    except:
        pass


def create_two_page_coa_pdf(english_text, russian_text):
    """Создает двухстраничный PDF-документ в памяти буфера"""
    buffer = io.BytesIO()
    # Настройка документа с полями по 40 поинтов
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    story = []

    # Страница 1: Английский оригинал
    build_pdf_page_elements(story, english_text, is_russian=False)
    
    # Принудительный разрыв страницы, чтобы перевод шел строго со 2 листа
    story.append(PageBreak())

    # Страница 2: Русский перевод
    build_pdf_page_elements(story, russian_text, is_russian=True)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# 5. ПРОВЕРКА КЛЮЧЕЙ SECRETS
# ==========================================
def check_secrets():
    status = {"yandex": False, "github": False}
    if "yandex" in st.secrets:
        y_sec = st.secrets["yandex"]
        if y_sec.get("folder_id") and y_sec.get("api_key"):
            status["yandex"] = True
    if "github" in st.secrets:
        g_sec = st.secrets["github"]
        if g_sec.get("token") and g_sec.get("repo") and g_sec.get("branch"):
            status["github"] = True
    return status

secrets_status = check_secrets()

# ==========================================
# 6. БОКОВАЯ ПАНЕЛЬ (SIDEBAR)
# ==========================================
with st.sidebar:
    st.header("⚙️ Настройки и Статус")
    st.subheader("Интеграции")
    if secrets_status["yandex"]:
        st.success("Yandex GPT API: Подключено")
    else:
        st.error("Yandex GPT API: Ключи не найдены")
    if secrets_status["github"]:
        st.success("GitHub Repository: Подключено")
    else:
        st.error("GitHub Repository: Данные не найдены")

# ==========================================
# 7. ИНТЕРФЕЙС ПРИЛОЖЕНИЯ
# ==========================================
st.title("🧪 Переводчик Сертификатов Анализа (CoA)")
st.write("Формирование двухстраничных PDF сертификатов: Лист 1 (ENG), Лист 2 (RUS).")

col_input, col_preview = st.columns([1, 1])
raw_text_to_translate = ""

with col_input:
    st.subheader("📥 Входные данные")
    tab_file, tab_text = st.tabs(["📄 Загрузить файл", "📝 Вставить текст"])
   
    with tab_text:
        input_text = st.text_area(
            "Вставьте скопированный текст сертификата:",
            height=300,
            placeholder="PRODUCT NAME | ADENOSINE\nCAS NUMBER | 58-61-7..."
        )
        if input_text:
            raw_text_to_translate = input_text
        
    with tab_file:
        uploaded_file = st.file_uploader(
            "Выберите файл сертификата (.docx, .pdf, .txt)", 
            type=["docx", "pdf", "txt"]
        )
        if uploaded_file is not None:
            try:
                raw_text_to_translate = extract_text_from_file(uploaded_file)
                st.success(f"Файл '{uploaded_file.name}' успешно распарсен!")
                with st.expander("🔍 Посмотреть результат парсинга"):
                    st.text(raw_text_to_translate[:1000] + "...")
            except Exception as e:
                st.error(f"Ошибка при парсинге файла: {e}")
   
    st.write("---")
    st.subheader("✏️ Корректировка данных")
    custom_product_name_ru = st.text_input(
        "Название продукта на русском (опционально):",
        placeholder="Например: АДЕНОЗИН, ХЧ (если оставить пустым — переведет автоматически)",
        help="Введенное сюда название гарантированно попадет в финальный русский PDF вместо автоматического перевода."
    )
    st.write("---")
    start_translation = st.button("🚀 Перевести и собрать PDF", type="primary", use_container_width=True)

with col_preview:
    st.subheader("📤 Результат перевода и сборки")
    
    if start_translation:
        if not raw_text_to_translate.strip():
            st.warning("Пожалуйста, предоставьте текст для перевода.")
        else:
            with st.spinner("Перевод текста и верстка двухстраничного PDF..."):
                # 1. Запускаем переводчик для получения русской версии
                russian_translated_text = process_coa_translation(raw_text_to_translate, custom_name_ru=custom_product_name_ru)
                
                # 2. Передаем английский оригинал и русский перевод в генератор PDF
                pdf_data = create_two_page_coa_pdf(raw_text_to_translate, russian_translated_text)
                
                st.success("Двухстраничный документ PDF успешно сформирован!")
                
                tab_view_text, tab_download = st.tabs(["👀 Русский перевод", "💾 Скачать PDF"])
                
                with tab_view_text:
                    st.text_area("Текст перевода для сверки:", value=russian_translated_text, height=350)
                    
                with tab_download:
                    st.write("Ваш двуязычный PDF-сертификат готов:")
                    st.download_button(
                        label="📥 Скачать двухстраничный PDF",
                        data=pdf_data,
                        file_name="CoA_Dual_Language.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
    else:
        st.info("Результат генерации PDF и кнопка скачивания появятся здесь после запуска.")

# ==========================================
# 8. ИНТЕРАКТИВНОЕ УПРАВЛЕНИЕ СЛОВАРЕМ
# ==========================================
st.write("---")
st.subheader("📖 Редактор словаря перевода (dictionary.csv)")

local_dict = st.session_state["coa_dictionary"]
if local_dict:
    import pandas as pd
    dict_data = {
        "english": list(local_dict.keys()),
        "russian": list(local_dict.values())
    }
    df_dict = pd.DataFrame(dict_data)
    
    st.write(f"Сейчас в словаре фраз: **{len(df_dict)}**. Вы можете изменять ячейки, добавлять новые строки внизу таблицы или удалять выделенные.")
   
    edited_df = st.data_editor(
        df_dict, 
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "english": st.column_config.TextColumn("Английская фраза (оригинал)", required=True),
            "russian": st.column_config.TextColumn("Русский перевод", required=True)
        }
    )
    
    # 3. Кнопка отправки изменений на GitHub
    if st.button("💾 Применить изменения и коммитить в GitHub", type="secondary"):
        with st.spinner("Отправляем обновленный словарь в репозиторий..."):
            # Проверяем, что текстовые поля не пустые, очищаем пробелы
            edited_df['english'] = edited_df['english'].astype(str).str.strip()
            edited_df['russian'] = edited_df['russian'].astype(str).str.strip()
            edited_df = edited_df[edited_df['english'] != ""]
            
            success = save_full_dictionary_to_github(edited_df)
            if success:
                st.success("Словарь успешно обновлен на GitHub и перезагружен в приложении!")
                st.rerun() # Перезапускаем приложение, чтобы обновить состояние памяти
else:
    st.warning("Словарь пуст или файл dictionary.csv не найден в корне проекта.")
